from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Word belgesini oluştur
doc = Document()

# Başlık
title = doc.add_heading('Burdur Bölgesi Büyükbaş İşletmeleri İçin Buzağı Bakım ve Kârlılık Raporu', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Alt Başlık
subtitle = doc.add_paragraph()
run = subtitle.add_run('Konu: CalfCare-Sim Veri Analizi ve Çiftlik Yönetim Tavsiyeleri\n')
run.bold = True
run.font.color.rgb = RGBColor(100, 100, 100)

# Giriş Metni
doc.add_paragraph('Kıymetli Yetiştiricimiz,\n\nHayvancılıkta yapılan ufak hataların bedeli maalesef yüksek veteriner masrafları veya can kayıpları ile ödeniyor. Sizlerin ahırda deneme-yanılma yaparak para ve hayvan kaybetmemesi için CalfCare-Sim adını verdiğimiz bir dijital ikiz sistemi geliştirdik.\n')

doc.add_heading('Bu Simülasyon Ne İşe Yarıyor?', level=2)
doc.add_paragraph('Bu sistem, ahırınızın dijital bir kopyasıdır. "Ağız sütünü 2 saat geç verirsem ne olur?", "Havalandırmayı kısarsam hayvan üşür mü yoksa amonyaktan mı zehirlenir?", "Açık ahırda kaç işçi çalıştırırsam hayvanı sağ salim büyütürüm?" gibi soruların cevaplarını, hayvanı gerçekten riske atmadan bilgisayar ortamında test etmemizi sağlar. İçerisine Burdur\'un iklim şartlarını, hastalıkların yayılma hızını ve günlük masrafları ekledik.')

doc.add_heading('150 Dijital Denemeden Ne Öğrendik?', level=2)
doc.add_paragraph('Sistemimizde 150 farklı senaryo denedik. Bu 150 buzağının 123\'ü sağlıklı büyüdü, maalesef 23\'ü bakım hataları yüzünden öldü, 4 senaryoda ise artan veteriner ve ilaç masrafları işletmeyi tamamen eksi bakiyeye düşürerek iflasa sürükledi.')

# Tablo 1: Kolostrum
doc.add_heading('Tablo 1: Ağız Sütü (Kolostrum) Zamanlamasının Etkisi', level=3)
table1 = doc.add_table(rows=1, cols=4)
table1.style = 'Light Shading Accent 1' # Word'ün şık mavi/gri teması
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Ağız Sütü Zamanı'
hdr_cells[1].text = 'Ort. Sağlık Skoru'
hdr_cells[2].text = 'Ort. Kârlılık'
hdr_cells[3].text = 'Risk Durumu'

records1 = [
    ('İlk 1 - 3 Saat İçinde', '%83', '18.500 TL', 'Güvenli'),
    ('10 - 11 Saat Gecikmeli', '%42', '10.300 TL', 'Riskli (Ölüm Artar)'),
    ('Hiç Verilmedi', '%13', '9.490 TL', 'Kritik (Ölüm)')
]
for a, b, c, d in records1:
    row_cells = table1.add_row().cells
    row_cells[0].text = a
    row_cells[1].text = b
    row_cells[2].text = c
    row_cells[3].text = d

doc.add_paragraph('\n')

# Tablo 2: Ahır Tipi
doc.add_heading('Tablo 2: Ahır Tipinin Etkisi', level=3)
table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Light Shading Accent 1'
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = 'Ahır Tipi'
hdr_cells2[1].text = 'Ort. Sağlık'
hdr_cells2[2].text = 'Risk Faktörü'

records2 = [
    ('İklim Kontrollü (Climate)', '%72', 'Elektrik/Maliyet'),
    ('Açık / Yarı Açık (Open)', '%55', 'Termal Şok/Zatürre')
]
for a, b, c in records2:
    row_cells = table2.add_row().cells
    row_cells[0].text = a
    row_cells[1].text = b
    row_cells[2].text = c

# Aksiyon Planı
doc.add_heading('Çiftlik Yönetimi İçin Aksiyon Planı', level=2)
doc.add_paragraph('1. İlk 2 Saat Kuralı: Doğum gece yarısı bile olsa ağız sütünü ilk 2 saat içinde verin. Hayvan doğal zırhını giyer, kârlılığınız %40 artar.', style='List Bullet')
doc.add_paragraph('2. Açık Ahırda Gözlem: İklimlendirmeli ahır yatırım maliyeti yüksekse ve açık ahır kullanıyorsanız, "Bakıcı Sayısını" artırın. Islanan hayvan anında fark edilmelidir.', style='List Bullet')
doc.add_paragraph('3. Havalandırma: "Hayvanlar üşümesin" diyerek havalandırmayı tamamen kapatırsanız amonyak gazı birikir. İdeal havalandırma solunum hastalıklarını bıçak gibi keser.', style='List Bullet')

doc.add_paragraph('\nÖzetle: Çiftlikte kârlılık, mucizevi ilaçlarla değil, doğumdan sonraki ilk 3 saatteki dikkatinizle sağlanmaktadır.\n')

doc.add_paragraph('Hazırlayan:\nAykut Çakıcı\nVeri Analisti ve Simülasyon Geliştirici')

# Dosyayı Kaydet
doc.save('CalfCare_Raporu.docx')
print("Kanki, Word dosyan başarıyla oluşturuldu! Bulunduğun klasöre bak.")